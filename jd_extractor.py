import os

try:
    import fitz  # PyMuPDF
    HAS_FITZ = True
except ImportError:
    HAS_FITZ = False

import PyPDF2
import docx


def extract_text_from_jd_pdf(file_obj):
    if HAS_FITZ:
        try:
            file_obj.seek(0)
            data = file_obj.read()
            doc = fitz.open(stream=data, filetype="pdf")
            text = "\n".join(page.get_text() for page in doc)
            doc.close()
            if text.strip():
                return text
        except Exception:
            pass  # fall through to PyPDF2

    file_obj.seek(0)
    reader = PyPDF2.PdfReader(file_obj)
    text = ""
    for page in reader.pages:
        extracted = page.extract_text()
        if extracted:
            text += extracted + "\n"
    return text


def extract_text_from_jd_docx(file_obj):
    file_obj.seek(0)
    document = docx.Document(file_obj)
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    # Also pull text out of tables, since some JDs put requirements in a table
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)
    return "\n".join(paragraphs)


def extract_jd_text(jd_text_field, jd_file):
    """
    Single entry point used by the Flask route. Accepts pasted text and/or
    an uploaded file — pasted text wins if both are somehow provided,
    since it's the more direct/explicit signal.

    Returns plain text, or "" if nothing usable was provided.
    """
    if jd_text_field and jd_text_field.strip():
        return jd_text_field.strip()

    if jd_file and jd_file.filename:
        ext = os.path.splitext(jd_file.filename)[1].lower()
        if ext == ".pdf":
            return extract_text_from_jd_pdf(jd_file).strip()
        elif ext in (".docx", ".doc"):
            if ext == ".doc":
                # python-docx cannot read legacy .doc (binary OLE format)
                raise ValueError(
                    "Legacy .doc files aren't supported — please upload as .docx or paste the JD text."
                )
            return extract_text_from_jd_docx(jd_file).strip()
        else:
            raise ValueError(f"Unsupported JD file type: {ext}. Use PDF, DOCX, or paste text.")

    return ""